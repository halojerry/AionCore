#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""DOCX 读取与结构理解（light-file-reading）。

依赖：python-docx。读段落/标题/run 样式、表格、节页边距、核心属性、页眉页脚。
注意：python-docx 不读修订(tracked changes)、不渲染。需要修订/批注请走
references/DOCX-REF.md 的 pandoc --track-changes=all 或裸 XML 路线。

CLI（处理真实文件）：
    python docx_read.py headings  file.docx
    python docx_read.py paragraphs file.docx
    python docx_read.py runs       file.docx
    python docx_read.py tables     file.docx
    python docx_read.py layout     file.docx
    python docx_read.py headers    file.docx
    python docx_read.py props      file.docx
自检（离线、自清理）：
    python docx_read.py --selftest
"""
import sys
sys.stdout.reconfigure(encoding="utf-8")  # Windows 控制台默认 GBK，强制 UTF-8 防乱码


def _docx():
    """惰性导入 python-docx，缺失时给 pip 提示而非裸 ImportError。"""
    try:
        import docx  # noqa: F401
        return docx
    except ImportError as e:  # pragma: no cover - 仅依赖缺失时触发
        raise SystemExit("缺少 python-docx，请先安装：pip install python-docx") from e


# 标题样式名：中英双语都认；脱离语言优先读 w:outlineLvl。
_HEADING_PREFIXES = ("Heading", "标题")          # Heading 1 / 标题 1
_TITLE_NAMES = ("Title", "题目", "标题")          # 文档主标题（level 0）
_LARGE_DOC_PARAS = 20000                          # 超此段数告警（可能太大）


def _iter_block_paragraphs(parent):
    """递归产出 parent 下所有段落，包括表格单元格内的段落。

    parent 可为 Document、_Cell。按文档先后顺序遍历正文 body 与表格嵌套段落。
    """
    from docx.document import Document as _Doc
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph

    if isinstance(parent, _Doc):
        elm = parent.element.body
    elif isinstance(parent, _Cell):
        elm = parent._tc
    else:  # pragma: no cover - 防御
        elm = parent.element

    for child in elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            tbl = Table(child, parent)
            for row in tbl.rows:
                for cell in row.cells:
                    yield from _iter_block_paragraphs(cell)


def _heading_level(p):
    """返回段落标题层级：w:outlineLvl 优先（0 起，+1 得 1..9），
    否则按 style 名（中英）解析。非标题返回 None；主标题返回 0。"""
    # 1) 直接读段落属性里的 outlineLvl（脱离语言、最可靠）
    pPr = p._p.pPr
    if pPr is not None:
        ol = pPr.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}outlineLvl"
        )
        if ol is not None:
            val = ol.get(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
            )
            if val is not None:
                try:
                    return int(val) + 1
                except ValueError:
                    pass
    # 2) 退回 style 名（中英双语）
    name = p.style.name if p.style else None
    if not name:
        return None
    # 先判主标题（精确匹配）：中文「标题」既是 Title 又是 Heading 前缀，
    # 必须先认精确的 Title，避免把 bare「标题」误判成 level 1。
    if name in _TITLE_NAMES:
        return 0
    for pre in _HEADING_PREFIXES:
        if name.startswith(pre):
            tail = name[len(pre):].strip()
            try:
                return int(tail) if tail else 1
            except ValueError:
                return 1
    return None


def read_paragraphs(path):
    """返回 [(style_name, text), ...]，跳过空段，含表格内段落。用于看章节骨架。"""
    docx = _docx()
    doc = docx.Document(path)
    out = []
    for p in _iter_block_paragraphs(doc):
        if p.text.strip():
            style = p.style.name if p.style else ""
            out.append((style, p.text))
    if len(out) > _LARGE_DOC_PARAS:
        print(f"[warn] 段落数 {len(out)} 超过 {_LARGE_DOC_PARAS}，文档较大、解析可能偏慢",
              file=sys.stderr)
    return out


def read_headings(path):
    """只取标题段 → [(level, text)]。level 优先 w:outlineLvl（脱离语言），
    否则按中英 style 名（Heading/标题/Title/题目）解析；主标题 level=0。
    含表格内的标题段。"""
    docx = _docx()
    doc = docx.Document(path)
    out = []
    for p in _iter_block_paragraphs(doc):
        if not p.text.strip():
            continue
        lvl = _heading_level(p)
        if lvl is not None:
            out.append((lvl, p.text))
    return out


def read_runs(path):
    """逐段逐 run 读样式（bold/italic/size/font），用于提取格式要求。
    返回 [{'para': i, 'text':.., 'bold':.., 'italic':.., 'size':.., 'font':..}]。"""
    docx = _docx()
    doc = docx.Document(path)
    rows = []
    for i, p in enumerate(doc.paragraphs):
        for r in p.runs:
            if not r.text:
                continue
            sz = r.font.size
            rows.append({
                "para": i,
                "text": r.text,
                "bold": r.bold,
                "italic": r.italic,
                "size": sz.pt if sz is not None else None,
                "font": r.font.name,
            })
    return rows


def read_tables(path):
    """读所有表格 → [list[list[str]]]（含表头行）。"""
    docx = _docx()
    doc = docx.Document(path)
    out = []
    for t in doc.tables:
        out.append([[c.text for c in row.cells] for row in t.rows])
    return out


def read_layout(path):
    """读节级页面格式（页边距/纸张），用于提取模板规范。返回 [dict]。"""
    docx = _docx()
    doc = docx.Document(path)

    def emu_in(v):
        return round(v.inches, 3) if v is not None else None  # Length→inch

    secs = []
    for s in doc.sections:
        secs.append({
            "page_w_in": emu_in(s.page_width),
            "page_h_in": emu_in(s.page_height),
            "margin_top_in": emu_in(s.top_margin),
            "margin_bottom_in": emu_in(s.bottom_margin),
            "margin_left_in": emu_in(s.left_margin),
            "margin_right_in": emu_in(s.right_margin),
        })
    return secs


def read_headers_footers(path):
    """读每节页眉/页脚文字 → [{'section':i,'header':str,'footer':str}]。
    空白页眉/页脚返回空串。"""
    docx = _docx()
    doc = docx.Document(path)

    def _text(container):
        if container is None:
            return ""
        parts = [p.text for p in container.paragraphs if p.text.strip()]
        return "\n".join(parts)

    out = []
    for i, s in enumerate(doc.sections):
        out.append({
            "section": i,
            "header": _text(s.header),
            "footer": _text(s.footer),
        })
    return out


def read_core_props(path):
    """读核心属性（作者/标题/创建/修改时间等），与 pdf_ops.read_meta 对齐。"""
    docx = _docx()
    doc = docx.Document(path)
    cp = doc.core_properties

    def _iso(dt):
        return dt.isoformat() if dt is not None else None

    return {
        "title": cp.title or None,
        "author": cp.author or None,
        "subject": cp.subject or None,
        "keywords": cp.keywords or None,
        "created": _iso(cp.created),
        "modified": _iso(cp.modified),
        "last_modified_by": cp.last_modified_by or None,
        "revision": cp.revision,
    }


def _selftest():
    """合成 docx（含中文标题、表格内标题、页眉页脚、核心属性）跑全流程，
    断言后清理临时目录，离线无残留。"""
    import os
    import tempfile
    docx = _docx()
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Pt
    from docx.oxml.ns import qn

    def _add_zh_heading_style(doc, name, base_outline):
        """注册一个中文名段落样式并写 outlineLvl（0 起），模拟中文模板。"""
        st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        pPr = st.element.get_or_add_pPr()
        ol = pPr.makeelement(qn("w:outlineLvl"), {qn("w:val"): str(base_outline)})
        pPr.append(ol)
        return st

    tmp = tempfile.mkdtemp(prefix="docxread_")
    try:
        src = os.path.join(tmp, "sample.docx")
        doc = Document()
        doc.core_properties.author = "光"
        doc.core_properties.title = "测试文档"
        doc.add_heading("Main Title", 0)
        doc.add_heading("Section One", 1)
        p = doc.add_paragraph()
        r = p.add_run("bold words")
        r.bold = True
        r.font.size = Pt(14)
        doc.add_heading("Section Two", 2)
        doc.add_paragraph("plain body text")
        # 中文标题 fixture：
        #  - 「标题 1」带 outlineLvl=0 → 经 outlineLvl 路径得 level 1（脱离语言）
        #  - 「标题 3」带 outlineLvl=2 → level 3，放表格内验证递归
        #  - 「标题」(Title) 无 outlineLvl → 经中文 style 名匹配得主标题 level 0
        _add_zh_heading_style(doc, "标题 1", 0)
        _add_zh_heading_style(doc, "标题 3", 2)
        zh_title_style = doc.styles.add_style("标题", WD_STYLE_TYPE.PARAGRAPH)
        zh1 = doc.add_paragraph("第一章 引言")
        zh1.style = "标题 1"
        zh_title = doc.add_paragraph("中文主标题")
        zh_title.style = zh_title_style
        t = doc.add_table(rows=2, cols=2)
        t.rows[0].cells[0].text = "H1"
        t.rows[0].cells[1].text = "H2"
        t.rows[1].cells[0].text = "a"
        t.rows[1].cells[1].text = "b"
        # 表格单元格内放一个标题段，验证递归遍历
        cell_p = t.rows[1].cells[1].add_paragraph("表内小节")
        cell_p.style = "标题 3"
        # 页眉/页脚
        sec = doc.sections[0]
        sec.header.paragraphs[0].text = "页眉文字"
        sec.footer.paragraphs[0].text = "第 X 页"
        doc.save(src)

        paras = read_paragraphs(src)
        assert any("plain body text" == txt for _, txt in paras), paras
        assert any("表内小节" == txt for _, txt in paras), "表格内段落未被遍历"

        heads = read_headings(src)
        assert (0, "Main Title") in heads and (1, "Section One") in heads, heads
        # X-1：中文 style 名标题必须被识别
        assert (1, "第一章 引言") in heads, f"中文标题 1 漏判: {heads}"
        assert (0, "中文主标题") in heads, f"中文主标题漏判: {heads}"
        # X-4：表格内标题段也应被收
        assert (3, "表内小节") in heads, f"表格内标题漏判: {heads}"

        runs = read_runs(src)
        assert any(rr["bold"] and rr["size"] == 14.0 for rr in runs), runs

        tables = read_tables(src)
        assert tables and tables[0][0] == ["H1", "H2"], tables

        layout = read_layout(src)
        assert layout and layout[0]["margin_top_in"] is not None, layout

        hf = read_headers_footers(src)
        assert hf and hf[0]["header"] == "页眉文字" and hf[0]["footer"] == "第 X 页", hf

        props = read_core_props(src)
        assert props["author"] == "光" and props["title"] == "测试文档", props
    finally:
        import shutil
        shutil.rmtree(tmp, ignore_errors=True)
    print("docx_read self-test OK: "
          "paragraphs/headings(中英+表内)/runs/tables/layout/headers/props all passed")


def _cli(argv):
    import argparse
    ap = argparse.ArgumentParser(
        prog="docx_read.py", description="DOCX 深度读取（章节/样式/表格/页面/页眉/属性）")
    ap.add_argument("--selftest", action="store_true",
                    help="跑合成自检（离线、自清理），CI 用此标志")
    sub = ap.add_subparsers(dest="cmd")
    for name, help_ in [
        ("headings", "列标题大纲 (level, text)"),
        ("paragraphs", "列非空段落 (style, text)，含表格内段落"),
        ("runs", "逐 run 样式 bold/italic/size/font"),
        ("tables", "所有表格行列文本"),
        ("layout", "节级页边距/纸张"),
        ("headers", "每节页眉/页脚文字"),
        ("props", "核心属性 作者/标题/时间"),
    ]:
        sp = sub.add_parser(name, help=help_)
        sp.add_argument("file", help="目标 .docx 路径")
    args = ap.parse_args(argv)

    if args.selftest:
        _selftest()
        return
    if not args.cmd:
        ap.print_help()
        raise SystemExit(2)

    import json
    fns = {
        "headings": read_headings,
        "paragraphs": read_paragraphs,
        "runs": read_runs,
        "tables": read_tables,
        "layout": read_layout,
        "headers": read_headers_footers,
        "props": read_core_props,
    }
    result = fns[args.cmd](args.file)
    print(json.dumps(result, ensure_ascii=False, indent=2, default=str))


if __name__ == "__main__":
    _cli(sys.argv[1:])
